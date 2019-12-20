'''
Reginald Zaccardi-Richey
Random Taco Cook Book Generator
'''

'''
github link - https://github.com/reginaldzacc/Final
'''

'''
This block of code is for importing the mods 
needed to handle my URL and create/edit word 
documents.
'''
import requests  # Imports requests mod to recieve data from api as JSON
import docx  # Imports docx for creating and editing word documents

'''
This block of code is variables needed to run the program
'''
document = docx.Document()  # this variable creates my empty word document for future use.
url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'  # this variable is the URL for the API used to get our JSON data

'''
This block of code is where the main() function is. 
This is where most of the code will be run.
this block of code will try to request JSON data 
using our URL, Pull the data from that request, and 
add a paragraphs and headers to word with the recipe. If there are 
errors the program will tell the user instead of crashing.
'''


def main(url):  # creates main() function
    try:  # this line 'trys' to run the following code
        data = requests.get(url).json()  # this line returns our API request as the variable 'data' using our URL variable

        seasoningName = (data["seasoning"]["name"])# creates seasoningName variable with information from our returned data
        condimentName = (data["condiment"]["name"])# creates condimentName variable with information from our returned data
        mixinName = (data["mixin"]["name"])# creates mixinName variable with information from our returned data
        base_layerName = (data["base_layer"]["name"])# creates base layerName variable with information from our returned data
        shellName = (data["shell"]["name"])# creates shellName variable with information from our returned data
        seasoning = (data["seasoning"]["recipe"])  # creates seasoning variable with information from our returned data
        condiment = (data["condiment"]["recipe"])  # creates condiment variable with information from our returned data
        mixin = (data["mixin"]["recipe"])  # creates mixin variable with information from our returned data
        base_layer = (data["base_layer"]["recipe"])  # creates base layer variable with information from our returned data
        shell = (data["shell"]["recipe"])  # creates shell variable with information from our returned data

        document.add_heading(f'{base_layerName} with {mixinName}, {seasoningName} and {condimentName} in {shellName}', 0)#prints the heading for each recipe to doc
        document.add_heading(f'{base_layerName}', 1)#prints base layer heading to doc
        document.add_paragraph(f'{base_layer}')#prints base layer recipe to doc
        document.add_heading(f'{mixinName}', 1)#prints mixin heading to doc
        document.add_paragraph(f'{mixin}')#prints mixin recipe to doc
        document.add_heading(f'{seasoningName}', 1)#prints seasoning heading to doc
        document.add_paragraph(f'{seasoning}')#prints seasoning recipe to doc
        document.add_heading(f'{condimentName}', 1)#prints condiment heading to doc
        document.add_paragraph(f'{condiment}')#prints condiment recipe to doc
        document.add_heading(f'{shellName}', 1)#prints shell heading to doc
        document.add_paragraph(f'{shell}')#prints shell recipe to doc
        document.add_page_break()#adds page break after recipe is done
    except ConnectionError:  # output if Connection Error
        print('Connection Error')  # what user sees if Connection Error
    except TimeoutError:  # output if Timeout error
        print('Timeout Error')  # what user sees if timeout Error
    except:  # broad error output
        print('Error')  # what user sees if Broad, missed Error


'''
this block of code will try to get input from 
user for amount of recipes they would like generated.
'''
while True:
    try:  # 'trys' to run following code
        recipes = int(input('How many recipes would you like in your Random cookbook?\nPlease enter a whole number in numeric form: '))# this line asks user for a whole number of recipes they would like
        break  # breaks loop
    except ValueError:  # error exception if wrong input data type
        print('Wrong Data Type. Try again')  # error output to user


'''
This block is meant to generate the tacobook using main looped as many times as the user wanted. 
'''
print('Generating TacoBook')#lets user know the tacobook is being generated
document.add_heading('Random Taco Cookbook', 0)#prints man heading to doc
document.add_picture('tacobookCover.jpeg')#adds cover photo to doc
document.add_heading('Credits', 1)#prints credits heading to doc
document.add_paragraph('Taco image : Photo by Xavier Crook on Unsplash', style='List Bullet')#prints bullet list to doc - photo credit
document.add_paragraph('Recipes from : https://taco-1150.herokuapp.com/random/?full_taco=true', style='List Bullet')#prints bullet list to doc - API credit
document.add_paragraph('Code by Reginald Zaccardi-Richey', style='List Bullet')#prints bullet list to doc - My credit
document.add_page_break()#adds page break after cover
for number in range(recipes):#For loop to create recipes
    main(url)#runs our main function
document.save('TacoBook.docx')#saves the document under file name "TacoBook.docx
print('Your Random Taco Cookbook has been created with the file name \'TacoBook.docx\'')#notifys user the tacobook is complete and the file name
